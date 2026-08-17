"""★ 화면·워드·노션의 «내용»이 같은지 대조한다.

정본: 확정/07_출력/3_기준/01_성공기준.md **P3** —
「화면 ↔ 워드 ↔ 노션 내용이 다른 건수 = **0건 고정**」
검사 방법도 정본이 정해 두었다 — 「같은 원본으로 세 형태를 만들고
**텍스트를 정규화해 비교**한다」.

★ 왜 이 시험이 필요한가
  화면과 워드를 «따로» 그리면, 한쪽만 고쳤을 때 조용히 갈린다.
  사용자는 화면을 보고 판단하는데 면접에는 워드를 들고 간다 — 둘이 다르면 사고다.

★ 노션까지 붙은 뒤에는 세 형태를 모두 만든다. 데모에는 작가 본문이 없으므로
  P-117·P-118 회귀는 표시용 문장과 출처가 든 합성 보고서로 따로 맞댄다.
"""

from __future__ import annotations

import html as html_lib
import io
import re
from dataclasses import replace
from typing import Any

import pytest
from docx import Document
from fastapi.testclient import TestClient

from src.core.constants import RAW_SOURCE_LABEL
from src.features.export_docx.logic import build_docx
from src.features.export_notion.logic import build_blocks
from src.features.pipeline.demo import DemoPipeline, available_companies
from src.features.pipeline.port import ReportTable, UserInput
from src.features.provenance.sources import Source, SourceKind
from src.features.storage import reports
from src.web import main as web_main

app = web_main.app

#: 이 글자들은 형태마다 달라도 «내용이 다른 것»이 아니다 (여백·기호·꾸밈).
_NOISE = re.compile(r"[\s·•\-–—\[\]()〔〕「」『』:：,，.]+")


def normalize(text: str) -> str:
    """형태 간 비교용으로 다듬는다 — 정본 「텍스트를 정규화해 비교」."""
    return _NOISE.sub("", text)


def report_of(company: str):
    pipe = DemoPipeline()
    item = next(c for c in available_companies() if c["company"] == company)
    user_input = UserInput(
        company=company, job=item["job"], region="", posting_text=""
    )
    result = pipe.run(user_input, pipe.find_company(user_input))
    assert result.report is not None
    return result.report


def docx_text(report) -> str:
    """워드 바이트를 다시 열어 «보이는 글자»를 전부 모은다."""
    document = Document(io.BytesIO(build_docx(report)))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def screen_text(company: str) -> str:
    """화면 HTML에서 태그를 걷어내고 «보이는 글자»만 남긴다."""
    item = next(c for c in available_companies() if c["company"] == company)
    with TestClient(app) as client:
        form = {
            "company": company,
            "job": item["job"],
            "region": "서울",
            "posting_text": "x",
        }
        confirm = client.post("/confirm", data=form)
        ref = re.search(r'name="ref" value="([^"]*)"', confirm.text).group(1)
        run = client.post(
            "/run", data={**form, "legal_name": company, "ref": ref},
            follow_redirects=False,
        )
        job_id = run.headers["location"].rsplit("/", 1)[-1]
        for _ in range(60):
            if client.get(f"/api/progress/{job_id}").json()["finished"]:
                break
        html = client.get(f"/result/{job_id}").text
    body = re.sub(r"(?s)<(script|style).*?</\1>", " ", html)
    # ★ HTML 기호를 되돌린다 — 화면의 `MD&amp;A`는 사실 `MD&A`다.
    #   안 되돌리면 «내용은 같은데 다르다»고 잘못 판정한다.
    return html_lib.unescape(re.sub(r"<[^>]+>", "\n", body))


def notion_text(report) -> str:
    """노션 블록의 표 자식까지 내려가 사람이 보게 될 글자를 전부 모은다."""

    def rich_text_text(items: list[dict[str, Any]]) -> str:
        return "".join(item["text"]["content"] for item in items)

    def block_text(block: dict[str, Any]) -> list[str]:
        kind = block["type"]
        payload = block[kind]
        parts: list[str] = []
        if "rich_text" in payload:
            parts.append(rich_text_text(payload["rich_text"]))
        for cell in payload.get("cells", []):
            parts.append(rich_text_text(cell))
        for child in payload.get("children", []):
            parts.extend(block_text(child))
        return parts

    return "\n".join(
        part
        for block in build_blocks(report)
        for part in block_text(block)
        if part
    )


def _writer_report():
    """작가 성공·실패·표를 함께 넣고 저장 왕복한 합성 보고서."""
    base = report_of(COMPANIES[0])
    section = next(s for s in base.sections if s.cell != "5" and s.lines)
    written_section = replace(
        section,
        lines=[
            ("첫 번째 근거 원문입니다.", "조각 1·사업보고서"),
            ("두 번째 근거 원문입니다.", "조각 2·뉴스"),
        ],
        prose_lines=[
            ("검증된 첫 번째 표시용 문장입니다.", "조각 1·사업보고서"),
            ("검증된 두 번째 표시용 문장입니다.", "조각 2·뉴스"),
        ],
        tables=[
            ReportTable(
                caption="표시용 글과 함께 남아야 하는 표",
                headers=["구분", "비중"],
                rows=[["제품", "60%"]],
                cite="조각 3·사업보고서",
                numeric=True,
            )
        ],
    )
    # ★ 작가나 검증이 실패한 칸은 `prose_lines`가 없다. 이때도 원문이
    #   세 형태의 본문에서 사라지지 않는지 같은 시험에서 맞댄다.
    fallback_section = replace(
        section,
        cell="2",
        title="작가 실패 시 원문 복귀",
        lines=[("작가 실패 뒤 본문으로 돌아온 근거 원문입니다.", "조각 3·사업보고서")],
        prose_lines=[],
        tables=[],
    )
    report = replace(
        base,
        sections=[written_section, fallback_section],
        citations=[
            Source(
                number=1,
                kind=SourceKind.FILING,
                label="사업보고서 사업의 내용",
                disclosed_at="2026-03-15",
                collected_at="2026-08-16",
            ),
            Source(
                number=2,
                kind=SourceKind.NEWS,
                label="회사의 최근 실행 사례",
                published_at="2026-08-10",
                domain="example.com",
            ),
            Source(
                number=3,
                kind=SourceKind.FILING,
                label="사업보고서 매출 구성",
                disclosed_at="2026-03-15",
                collected_at="2026-08-16",
            ),
        ],
        generated_at="2026-08-16",
    )
    return reports.report_from_json(reports.report_to_json(report))


def _screen_for_report(report, monkeypatch: pytest.MonkeyPatch) -> tuple[str, str]:
    job_id = "writer-three-forms-p117"
    web_main._JOBS.pop(job_id, None)
    monkeypatch.setattr(web_main, "_load_saved_report", lambda _report_id: report)
    with TestClient(app) as client:
        html = client.get(f"/result/{job_id}").text
    body = re.sub(r"(?s)<(script|style).*?</\1>", " ", html)
    visible = html_lib.unescape(re.sub(r"<[^>]+>", "\n", body))
    return html, visible


#: 보고서가 나오는 데모 회사 중 골고루 (표 있는 곳·6·7·8 있는 곳 포함)
COMPANIES = ["글로벌머니익스프레스", "로보스타", "파마리서치", "넥스트증권"]


@pytest.mark.parametrize("company", COMPANIES)
def test_본문_문장이_워드에_하나도_빠짐없이_들어간다(company):
    """★ 화면에 있는 사실이 워드에 없으면 P3 위반이다."""
    report = report_of(company)
    made = normalize(docx_text(report))
    missing = [
        text
        for section in report.sections
        for text, _cite in section.lines
        if normalize(text) not in made
    ]
    assert not missing, f"워드에 빠진 문장 {len(missing)}개: {missing[:2]}"


@pytest.mark.parametrize("company", COMPANIES)
def test_빈칸_사유가_워드에도_들어간다(company):
    """빈칸 사유를 빠뜨리면 「회사에 자료가 없다」로 오해하게 된다."""
    report = report_of(company)
    made = normalize(docx_text(report))
    missing = [
        s.cell
        for s in report.sections
        if not s.is_filled and s.empty_reason and normalize(s.empty_reason) not in made
    ]
    assert not missing, f"사유가 빠진 칸: {missing}"


@pytest.mark.parametrize("company", COMPANIES)
def test_요구역량이_원문_그대로_워드에_들어간다(company):
    """5번은 «다듬지 않는다»가 규칙이다 (출력틀 규칙⑤)."""
    report = report_of(company)
    made = normalize(docx_text(report))
    missing = [r for r in report.requirements if normalize(r) not in made]
    assert not missing, f"워드에 빠진 요구역량 {len(missing)}개: {missing[:2]}"


@pytest.mark.parametrize("company", COMPANIES)
def test_표의_숫자가_워드에도_들어간다(company):
    """재무 표를 글자로 뭉개거나 빠뜨리면 안 된다."""
    report = report_of(company)
    made = normalize(docx_text(report))
    missing = [
        cell
        for section in report.sections
        for table in section.tables
        for row in table.rows
        for cell in row
        if cell and normalize(cell) not in made
    ]
    assert not missing, f"워드에 빠진 표 칸 {len(missing)}개: {missing[:3]}"


@pytest.mark.parametrize("company", COMPANIES)
def test_출처가_화면과_워드_양쪽에_있다(company):
    """★ 출처가 한쪽에만 있으면 P3 위반이다."""
    report = report_of(company)
    if not report.citations:
        pytest.skip("이 회사는 출처 재료가 없다 (데모 한계 — 문제로그 P-24)")
    made = normalize(docx_text(report))
    shown = normalize(screen_text(company))
    for source in report.citations:
        assert normalize(source.label) in made, f"워드에 없는 출처: {source.label}"
        assert normalize(source.label) in shown, f"화면에 없는 출처: {source.label}"


@pytest.mark.parametrize("company", COMPANIES)
def test_화면에_보이는_본문이_워드에도_있다(company):
    """양쪽을 «실제로» 만들어 맞대어 본다 — 정본이 정한 검사 방법."""
    report = report_of(company)
    shown = normalize(screen_text(company))
    made = normalize(docx_text(report))
    for section in report.sections:
        for text, _cite in section.lines:
            key = normalize(text)
            if key in shown:
                assert key in made, f"화면엔 있는데 워드엔 없다: {text[:40]}"


def test_검증된_본문_출처_원문이_화면_워드_노션에_모두_있다(
    monkeypatch: pytest.MonkeyPatch,
):
    """P-117·P-118·P-127 — 세 형태가 같은 실제 출처 번호를 내는지 맞댄다."""
    report = _writer_report()
    html, shown = _screen_for_report(report, monkeypatch)
    made = docx_text(report)
    notion = notion_text(report)

    표시용칸 = report.sections[0]
    복귀칸 = report.sections[1]
    for text, _cite in 표시용칸.prose_lines + 표시용칸.lines + 복귀칸.lines:
        assert normalize(text) in normalize(shown)
        assert normalize(text) in normalize(made)
        assert normalize(text) in normalize(notion)

    assert RAW_SOURCE_LABEL in shown
    assert RAW_SOURCE_LABEL in made
    assert RAW_SOURCE_LABEL in notion
    assert 'href="#src1"' in html
    assert 'href="#src2"' in html
    assert 'title="출처 1번"' in html and 'title="출처 2번"' in html
    assert "검증된 첫 번째 표시용 문장입니다. 〔1〕" in made
    assert "검증된 두 번째 표시용 문장입니다. 〔2〕" in made
    assert "검증된 첫 번째 표시용 문장입니다. 〔1〕" in notion
    assert "검증된 두 번째 표시용 문장입니다. 〔2〕" in notion
    for rendered in (made, notion):
        assert "조각 1·사업보고서" not in rendered
        assert "조각 2·뉴스" not in rendered

    표글자 = ("표시용 글과 함께 남아야 하는 표", "구분", "비중", "제품", "60%")
    for text in 표글자:
        assert normalize(text) in normalize(shown)
        assert normalize(text) in normalize(made)
        assert normalize(text) in normalize(notion)

    # 세 형태 모두 표시용 글 → 원문 보기 → 근거 원문 → 표 순서다.
    순서표 = (
        "검증된 첫 번째 표시용 문장입니다.",
        RAW_SOURCE_LABEL,
        "첫 번째 근거 원문입니다.",
        "표시용 글과 함께 남아야 하는 표",
    )
    for rendered in (shown, made, notion):
        positions = [normalize(rendered).index(normalize(text)) for text in 순서표]
        assert positions == sorted(positions)
