"""Word 보고서의 시각 스타일과 저수준 OOXML 설정.

이 모듈은 보고서 내용을 만들지 않는다. ``logic``이 조립한 문단과 표에
``standard_business_brief`` 스타일, 페이지 장식, 안전한 메타데이터 및
OOXML 구조만 적용한다.
"""

from __future__ import annotations

from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from src.features.export_docx.constants import (
    BODY_LINE_SPACING,
    BODY_SPACE_AFTER_PT,
    BULLET_HANGING_DXA,
    BULLET_LINE_SPACING,
    BULLET_MARKER_DXA,
    BULLET_SPACE_AFTER_PT,
    BULLET_STYLE,
    BULLET_TEXT_DXA,
    COLOR_ACCENT_DARK_RGB,
    COLOR_ACCENT_RGB,
    COLOR_INK_RGB,
    COLOR_MUTED_RGB,
    COLOR_TABLE_BORDER_HEX,
    FONT_NAME,
    FONT_SIZE_BODY_PT,
    FONT_SIZE_HEADING_1_PT,
    FONT_SIZE_HEADING_3_PT,
    FONT_SIZE_HEADING_PT,
    FONT_SIZE_MUTED_PT,
    FONT_SIZE_SUBTITLE_PT,
    FONT_SIZE_TITLE_PT,
    HEADER_FOOTER_DISTANCE_IN,
    HEADING_1_SPACE_AFTER_PT,
    HEADING_1_SPACE_BEFORE_PT,
    HEADING_2_SPACE_AFTER_PT,
    HEADING_2_SPACE_BEFORE_PT,
    HEADING_3_SPACE_AFTER_PT,
    HEADING_3_SPACE_BEFORE_PT,
    PAGE_HEIGHT_IN,
    PAGE_MARGIN_IN,
    PAGE_WIDTH_IN,
    TABLE_CELL_BOTTOM_DXA,
    TABLE_CELL_END_DXA,
    TABLE_CELL_START_DXA,
    TABLE_CELL_TOP_DXA,
    TABLE_INDENT_DXA,
    TABLE_WIDTH_DXA,
)
from src.features.pipeline.port import ReportTable


def _set_font(
    run: Run,
    *,
    size_pt: float = FONT_SIZE_BODY_PT,
    bold: bool = False,
    color_rgb: tuple[int, int, int] | None = None,
) -> None:
    """실행(run) 하나에 로마자·동아시아 서체를 모두 지정한다."""
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_rgb is not None:
        run.font.color.rgb = RGBColor(*color_rgb)
    r_pr = run.font.element.get_or_add_rPr()
    fonts = r_pr.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), FONT_NAME)


def _set_style_font(
    style,
    *,
    size_pt: float,
    color_rgb: tuple[int, int, int],
    bold: bool = False,
) -> None:
    """스타일에도 로마자·한글 글꼴과 색을 명시해 기본값에 기대지 않는다."""
    style.font.name = FONT_NAME
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(*color_rgb)
    fonts = style.font.element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), FONT_NAME)


def _set_paragraph_rhythm(
    style,
    *,
    before_pt: float,
    after_pt: float,
    line_spacing: float,
    keep_with_next: bool = False,
) -> None:
    paragraph = style.paragraph_format
    paragraph.space_before = Pt(before_pt)
    paragraph.space_after = Pt(after_pt)
    paragraph.line_spacing = line_spacing
    paragraph.keep_with_next = keep_with_next
    paragraph.widow_control = True


def _paragraph_style(doc: DocxDocument, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _configure_styles(doc: DocxDocument) -> None:
    """``standard_business_brief`` 토큰을 실제 Word 스타일에 기록한다."""
    normal = doc.styles["Normal"]
    _set_style_font(
        normal,
        size_pt=FONT_SIZE_BODY_PT,
        color_rgb=COLOR_INK_RGB,
    )
    _set_paragraph_rhythm(
        normal,
        before_pt=0,
        after_pt=BODY_SPACE_AFTER_PT,
        line_spacing=BODY_LINE_SPACING,
    )

    title = doc.styles["Title"]
    _set_style_font(
        title,
        size_pt=FONT_SIZE_TITLE_PT,
        color_rgb=COLOR_INK_RGB,
        bold=True,
    )
    _set_paragraph_rhythm(
        title,
        before_pt=0,
        after_pt=6,
        line_spacing=1.05,
        keep_with_next=True,
    )
    subtitle = doc.styles["Subtitle"]
    _set_style_font(
        subtitle,
        size_pt=FONT_SIZE_SUBTITLE_PT,
        color_rgb=COLOR_MUTED_RGB,
    )
    _set_paragraph_rhythm(
        subtitle,
        before_pt=0,
        after_pt=14,
        line_spacing=BODY_LINE_SPACING,
        keep_with_next=True,
    )

    heading_tokens = (
        (
            "Heading 1",
            FONT_SIZE_HEADING_1_PT,
            COLOR_ACCENT_RGB,
            HEADING_1_SPACE_BEFORE_PT,
            HEADING_1_SPACE_AFTER_PT,
        ),
        (
            "Heading 2",
            FONT_SIZE_HEADING_PT,
            COLOR_ACCENT_RGB,
            HEADING_2_SPACE_BEFORE_PT,
            HEADING_2_SPACE_AFTER_PT,
        ),
        (
            "Heading 3",
            FONT_SIZE_HEADING_3_PT,
            COLOR_ACCENT_DARK_RGB,
            HEADING_3_SPACE_BEFORE_PT,
            HEADING_3_SPACE_AFTER_PT,
        ),
    )
    for name, size, color, before, after in heading_tokens:
        style = doc.styles[name]
        _set_style_font(style, size_pt=size, color_rgb=color, bold=True)
        _set_paragraph_rhythm(
            style,
            before_pt=before,
            after_pt=after,
            line_spacing=BODY_LINE_SPACING,
            keep_with_next=True,
        )

    bullet = doc.styles[BULLET_STYLE]
    _set_style_font(
        bullet,
        size_pt=FONT_SIZE_BODY_PT,
        color_rgb=COLOR_INK_RGB,
    )
    _set_paragraph_rhythm(
        bullet,
        before_pt=0,
        after_pt=BULLET_SPACE_AFTER_PT,
        line_spacing=BULLET_LINE_SPACING,
    )
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)

    source = _paragraph_style(doc, "Report Source")
    _set_style_font(
        source,
        size_pt=FONT_SIZE_MUTED_PT,
        color_rgb=COLOR_MUTED_RGB,
    )
    _set_paragraph_rhythm(
        source,
        before_pt=0,
        after_pt=4,
        line_spacing=BODY_LINE_SPACING,
    )
    source.paragraph_format.left_indent = Inches(0.15)
    source.paragraph_format.first_line_indent = Inches(-0.15)

    caption = _paragraph_style(doc, "Report Table Caption")
    _set_style_font(
        caption,
        size_pt=FONT_SIZE_BODY_PT,
        color_rgb=COLOR_INK_RGB,
        bold=True,
    )
    _set_paragraph_rhythm(
        caption,
        before_pt=4,
        after_pt=4,
        line_spacing=BODY_LINE_SPACING,
        keep_with_next=True,
    )


# ECMA-376 CT_PPr/CT_TblPr sequence. Word가 복구 모드로 열지 않도록
# 공개 API가 없는 속성도 이 순서에 맞춰 삽입한다.
PPR_CHILD_ORDER: tuple[str, ...] = (
    "pStyle",
    "keepNext",
    "keepLines",
    "pageBreakBefore",
    "framePr",
    "widowControl",
    "numPr",
    "suppressLineNumbers",
    "pBdr",
    "shd",
    "tabs",
    "suppressAutoHyphens",
    "kinsoku",
    "wordWrap",
    "overflowPunct",
    "topLinePunct",
    "autoSpaceDE",
    "autoSpaceDN",
    "bidi",
    "adjustRightInd",
    "snapToGrid",
    "spacing",
    "ind",
    "contextualSpacing",
    "mirrorIndents",
    "suppressOverlap",
    "jc",
    "textDirection",
    "textAlignment",
    "textboxTightWrap",
    "outlineLvl",
    "divId",
    "cnfStyle",
    "rPr",
    "sectPr",
    "pPrChange",
)

TBLPR_CHILD_ORDER: tuple[str, ...] = (
    "tblStyle",
    "tblpPr",
    "tblOverlap",
    "bidiVisual",
    "tblStyleRowBandSize",
    "tblStyleColBandSize",
    "tblW",
    "jc",
    "tblCellSpacing",
    "tblInd",
    "tblBorders",
    "shd",
    "tblLayout",
    "tblCellMar",
    "tblLook",
    "tblCaption",
    "tblDescription",
    "tblPrChange",
)


def _set_child(parent, tag: str, **attributes: str):
    child = parent.find(qn(f"w:{tag}"))
    if child is None:
        child = OxmlElement(f"w:{tag}")
        parent.append(child)
    for name, value in attributes.items():
        child.set(qn(f"w:{name}"), value)
    return child


def _set_ordered_child(
    parent,
    tag: str,
    order: tuple[str, ...],
    **attributes: str,
):
    """자식을 만들거나 갱신하고 지정된 OOXML 스키마 순서로 다시 놓는다."""
    child = _set_child(parent, tag, **attributes)
    parent.remove(child)
    rank = order.index(tag)
    for index, sibling in enumerate(parent):
        sibling_name = sibling.tag.rsplit("}", 1)[-1]
        if sibling_name in order and order.index(sibling_name) > rank:
            parent.insert(index, child)
            break
    else:
        parent.append(child)
    return child


def _configure_bullet_numbering(doc: DocxDocument) -> None:
    """List Bullet의 numbering 정의도 0.25/0.5인치에 맞춘다."""
    style = doc.styles[BULLET_STYLE]
    num_pr = style.element.pPr.numPr
    if num_pr is None or num_pr.numId is None:
        raise ValueError("List Bullet 스타일에 실제 numbering 정의가 없습니다")
    num_id = str(num_pr.numId.val)
    numbering = doc.part.numbering_part.element
    nums = numbering.xpath(f"./w:num[@w:numId='{num_id}']")
    if not nums:
        raise ValueError("List Bullet numId를 numbering.xml에서 찾지 못했습니다")
    abstract_id = str(nums[0].abstractNumId.val)
    levels = numbering.xpath(
        f"./w:abstractNum[@w:abstractNumId='{abstract_id}']/w:lvl[@w:ilvl='0']"
    )
    if not levels:
        raise ValueError("List Bullet 0단계 정의를 찾지 못했습니다")
    p_pr = levels[0].find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        levels[0].append(p_pr)
    indent = _set_ordered_child(
        p_pr,
        "ind",
        PPR_CHILD_ORDER,
        left=str(BULLET_TEXT_DXA),
        hanging=str(BULLET_HANGING_DXA),
    )
    indent.attrib.pop(qn("w:firstLine"), None)
    tabs = _set_ordered_child(p_pr, "tabs", PPR_CHILD_ORDER)
    tab = _set_child(tabs, "tab", val="num", pos=str(BULLET_MARKER_DXA))
    tab.set(qn("w:leader"), "none")


def _set_page_geometry(doc: DocxDocument) -> None:
    """Letter 세로·1인치 여백·0.492인치 머리말 거리를 명시한다."""
    for section in doc.sections:
        section.page_width = Inches(PAGE_WIDTH_IN)
        section.page_height = Inches(PAGE_HEIGHT_IN)
        section.left_margin = Inches(PAGE_MARGIN_IN)
        section.right_margin = Inches(PAGE_MARGIN_IN)
        section.top_margin = Inches(PAGE_MARGIN_IN)
        section.bottom_margin = Inches(PAGE_MARGIN_IN)
        section.header_distance = Inches(HEADER_FOOTER_DISTANCE_IN)
        section.footer_distance = Inches(HEADER_FOOTER_DISTANCE_IN)


def _clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):  # noqa: SLF001 — 헤더 기본 문단 재사용
        paragraph._p.remove(child)  # noqa: SLF001


def _add_page_field(paragraph: Paragraph) -> None:
    run = paragraph.add_run()
    _set_font(run, size_pt=FONT_SIZE_MUTED_PT, color_rgb=COLOR_MUTED_RGB)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, placeholder, end):
        run._r.append(element)  # noqa: SLF001 — Word PAGE 필드는 공개 API가 없다


def _add_page_furniture(doc: DocxDocument, *, report_title: str) -> None:
    """표지는 비우고 본문 페이지에만 조용한 머리말과 쪽수를 둔다."""
    for section in doc.sections:
        section.different_first_page_header_footer = True

        first_header = section.first_page_header.paragraphs[0]
        _clear_paragraph(first_header)
        first_footer = section.first_page_footer.paragraphs[0]
        _clear_paragraph(first_footer)

        header = section.header.paragraphs[0]
        _clear_paragraph(header)
        header.paragraph_format.space_before = Pt(0)
        header.paragraph_format.space_after = Pt(0)
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header.add_run(report_title)
        _set_font(
            header_run,
            size_pt=FONT_SIZE_MUTED_PT,
            color_rgb=COLOR_MUTED_RGB,
        )

        footer = section.footer.paragraphs[0]
        _clear_paragraph(footer)
        footer.paragraph_format.space_before = Pt(0)
        footer.paragraph_format.space_after = Pt(0)
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        label = footer.add_run("페이지 ")
        _set_font(label, size_pt=FONT_SIZE_MUTED_PT, color_rgb=COLOR_MUTED_RGB)
        _add_page_field(footer)


def _clear_personal_metadata(doc: DocxDocument, *, title: str) -> None:
    """개인 정보는 비우되 보조기술에 필요한 문서 제목은 남긴다."""
    properties = doc.core_properties
    properties.author = ""
    properties.last_modified_by = ""
    properties.comments = ""
    properties.category = ""
    properties.keywords = ""
    properties.subject = ""
    properties.title = title


def _clear_revision_ids(doc: DocxDocument) -> None:
    """본문·머리말·꼬리말의 Word 편집 세션 식별값을 제거한다."""
    roots = [doc.element]
    for section in doc.sections:
        roots.extend(  # noqa: SLF001
            (
                section.header._element,
                section.footer._element,
                section.first_page_header._element,
                section.first_page_footer._element,
                section.even_page_header._element,
                section.even_page_footer._element,
            )
        )
    rsid_prefix = qn("w:rsidR").rsplit("}", 1)[0] + "}rsid"
    visited: set[int] = set()
    for root in roots:
        if id(root) in visited:
            continue
        visited.add(id(root))
        for element in root.iter():
            for attribute in list(element.attrib):
                if attribute.startswith(rsid_prefix):
                    del element.attrib[attribute]


def _shade_cell(cell: _Cell, hex_color: str) -> None:
    """표 칸 배경을 칠한다."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)  # noqa: SLF001


def _style_callout(
    paragraph: Paragraph,
    *,
    fill_hex: str,
    border_hex: str,
) -> None:
    """표 대신 문단 자체를 절제된 안내 상자로 만든다."""
    p_pr = paragraph._p.get_or_add_pPr()  # noqa: SLF001
    borders = _set_ordered_child(p_pr, "pBdr", PPR_CHILD_ORDER)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border_hex)
    borders.append(left)

    _set_ordered_child(
        p_pr,
        "shd",
        PPR_CHILD_ORDER,
        val="clear",
        color="auto",
        fill=fill_hex,
    )
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING


def _column_widths_dxa(table: ReportTable) -> tuple[int, ...]:
    """내용에 맞춰 열 폭을 나누되 합은 항상 9360 DXA로 맞춘다."""
    count = len(table.headers)
    if count <= 1:
        return (TABLE_WIDTH_DXA,)
    if table.numeric:
        first = round(TABLE_WIDTH_DXA * 0.32)
        tail = TABLE_WIDTH_DXA - first
        base = tail // (count - 1)
        widths = [first, *([base] * (count - 1))]
        widths[-1] += TABLE_WIDTH_DXA - sum(widths)
        return tuple(widths)

    columns = [
        [header, *(row[index] if index < len(row) else "" for row in table.rows)]
        for index, header in enumerate(table.headers)
    ]
    weights = [
        max(4, min(36, max(len(str(value)) for value in column)))
        for column in columns
    ]
    minimum = min(1100, TABLE_WIDTH_DXA // count)
    remainder = TABLE_WIDTH_DXA - minimum * count
    total_weight = sum(weights)
    widths = [minimum + (remainder * weight // total_weight) for weight in weights]
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return tuple(widths)


def _set_table_geometry(
    table: Table,
    widths_dxa: tuple[int, ...],
    *,
    repeat_header: bool = True,
) -> None:
    """모든 렌더러에서 같은 폭이 되도록 표 geometry를 고정 DXA로 기록한다."""
    if len(widths_dxa) != len(table.columns) or sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError("표 열 폭은 열 개수와 9360 DXA 합계를 정확히 맞춰야 합니다")

    table.autofit = False
    tbl_pr = table._tbl.tblPr  # noqa: SLF001
    _set_ordered_child(
        tbl_pr,
        "tblW",
        TBLPR_CHILD_ORDER,
        type="dxa",
        w=str(TABLE_WIDTH_DXA),
    )
    _set_ordered_child(
        tbl_pr,
        "tblInd",
        TBLPR_CHILD_ORDER,
        type="dxa",
        w=str(TABLE_INDENT_DXA),
    )
    _set_ordered_child(tbl_pr, "tblLayout", TBLPR_CHILD_ORDER, type="fixed")

    margins = _set_ordered_child(tbl_pr, "tblCellMar", TBLPR_CHILD_ORDER)
    for side, width in (
        ("top", TABLE_CELL_TOP_DXA),
        ("start", TABLE_CELL_START_DXA),
        ("bottom", TABLE_CELL_BOTTOM_DXA),
        ("end", TABLE_CELL_END_DXA),
    ):
        _set_child(margins, side, type="dxa", w=str(width))

    borders = _set_ordered_child(tbl_pr, "tblBorders", TBLPR_CHILD_ORDER)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        _set_child(
            borders,
            edge,
            val="single",
            sz="4",
            space="0",
            color=COLOR_TABLE_BORDER_HEX,
        )

    grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))  # noqa: SLF001
    for grid_column, width in zip(grid_columns, widths_dxa, strict=True):
        grid_column.set(qn("w:w"), str(width))

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa, strict=True):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()  # noqa: SLF001
            tc_width.set(qn("w:type"), "dxa")
            tc_width.set(qn("w:w"), str(width))
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING

    if repeat_header and table.rows:
        header_properties = table.rows[0]._tr.get_or_add_trPr()  # noqa: SLF001
        if header_properties.find(qn("w:tblHeader")) is None:
            repeating = OxmlElement("w:tblHeader")
            repeating.set(qn("w:val"), "true")
            header_properties.append(repeating)
