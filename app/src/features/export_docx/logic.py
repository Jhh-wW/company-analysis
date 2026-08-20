"""검증된 ``Report``를 편집 가능한 Word 보고서 바이트로 바꾼다.

canonical 보고서는 공통 출고 게이트를 먼저 통과시킨 뒤 화면·PDF·Notion과
같은 표지, 핵심 요약, 의미 기반 장 순서와 최종 출처표만 낸다. ``prose_lines``가
있으면 같은 사실의 근거 원문 ``lines``는 반복하지 않으며, 직무·수집 과정·완성도
같은 서비스 메타데이터는 문서에 넣지 않는다.
"""

from __future__ import annotations

import datetime as dt
import io
import urllib.parse

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from src.core.citations import citation_marker
from src.core.constants import section_display_heading
from src.features.export_docx.constants import (
    CITATIONS_NOTE,
    COLOR_INK_RGB,
    COLOR_MUTED_RGB,
    COLOR_TABLE_HEADER_HEX,
    CONTENT_TYPE_DOCX,
    DOCX_SUFFIX,
    FILENAME_ASCII_ALLOWED,
    FILENAME_ASCII_FALLBACK,
    FILENAME_ASCII_MIN_STEM,
    FILENAME_FALLBACK,
    FILENAME_FORBIDDEN_CHARS,
    FONT_SIZE_HEADING_PT,
    FONT_SIZE_MUTED_PT,
    FONT_SIZE_SUBTITLE_PT,
    FONT_SIZE_TABLE_PT,
    FONT_SIZE_TITLE_PT,
    HEADING_SOURCES,
    SOURCES_TABLE_WIDTHS_DXA,
    SUMMARY_TABLE_WIDTHS_DXA,
    TABLE_STYLE,
)
from src.features.export_docx.document_style import (
    _add_page_furniture,
    _clear_personal_metadata,
    _clear_revision_ids,
    _column_widths_dxa,
    _configure_bullet_numbering,
    _configure_styles,
    _set_font,
    _set_page_geometry,
    _set_table_geometry,
    _shade_cell,
)
from src.features.pipeline.port import (
    Report,
    ReportSection,
    ReportTable,
)
from src.features.provenance.sources import Source
from src.features.report_standard import (
    SECTION_BY_ID,
    build_published_report,
)

# ══════════════════════════════════════════════════════════
# 문단 헬퍼
# ══════════════════════════════════════════════════════════


def _add_heading(doc: DocxDocument, text: str, *, level: int) -> None:
    """제목 문단을 추가한다. `level=0`은 회사명(문서 제목), 그 외는 항목 제목."""
    heading = doc.add_heading(text, level=level)
    size = FONT_SIZE_TITLE_PT if level == 0 else FONT_SIZE_HEADING_PT
    heading.paragraph_format.keep_with_next = True
    for run in heading.runs:
        _set_font(run, size_pt=size, bold=True)


def _add_muted_paragraph(doc: DocxDocument, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    _set_font(run, size_pt=FONT_SIZE_MUTED_PT, color_rgb=COLOR_MUTED_RGB)


def _add_cited_prose(doc: DocxDocument, lines: list[tuple[str, str]]) -> None:
    """검증된 표시용 문장을 한 문단으로 잇고, 각 문장 뒤에 실제 출처를 붙인다."""
    paragraph = doc.add_paragraph()
    for index, (text, cite) in enumerate(lines):
        if index:
            spacer = paragraph.add_run(" ")
            _set_font(spacer)
        run = paragraph.add_run(text)
        _set_font(run)
        marker = citation_marker(cite)
        if marker:
            cite_run = paragraph.add_run(f" {marker}")
            _set_font(cite_run, size_pt=FONT_SIZE_MUTED_PT, color_rgb=COLOR_MUTED_RGB)


# ══════════════════════════════════════════════════════════
# 표
# ══════════════════════════════════════════════════════════


def _add_report_table(doc: DocxDocument, table: ReportTable) -> None:
    """숫자·글자 표 하나 — 문장으로 뭉개지 않고 워드 표 그대로 넣는다 (D13)."""
    caption = doc.add_paragraph(style="Report Table Caption")
    caption_run = caption.add_run(table.caption)
    _set_font(caption_run, bold=True)
    marker = citation_marker(table.cite)
    if marker:
        cite_run = caption.add_run(f" {marker}")
        _set_font(cite_run, size_pt=FONT_SIZE_MUTED_PT, color_rgb=COLOR_MUTED_RGB)

    n_cols = len(table.headers)
    docx_table = doc.add_table(rows=1, cols=n_cols)
    docx_table.style = TABLE_STYLE

    for idx, head in enumerate(table.headers):
        cell = docx_table.rows[0].cells[idx]
        run = cell.paragraphs[0].add_run(head)
        _set_font(run, bold=True, size_pt=FONT_SIZE_TABLE_PT)
        _shade_cell(cell, COLOR_TABLE_HEADER_HEX)
        if table.numeric and idx > 0:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for row in table.rows:
        cells = docx_table.add_row().cells
        for idx in range(n_cols):
            # ★ 방어적 — 정상 데이터라면 `ReportTable.is_valid`가 열 개수를
            #   보장하지만, 만에 하나 어긋나도 화면처럼 예외로 죽지 않는다.
            value = row[idx] if idx < len(row) else ""
            run = cells[idx].paragraphs[0].add_run(value)
            _set_font(run, size_pt=FONT_SIZE_TABLE_PT)
            if table.numeric and idx > 0:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_table_geometry(docx_table, _column_widths_dxa(table))


def _source_label(source: Source) -> str:
    label = (source.title or source.label).strip()
    publisher = source.publisher.strip()
    if publisher and publisher.casefold() not in label.casefold():
        return f"{label} · {publisher}"
    return label


def _source_status(source: Source) -> str:
    parts: list[str] = []
    if source.published_at:
        parts.append(f"{source.published_at} 보도")
    elif source.disclosed_at:
        parts.append(f"{source.disclosed_at} 공시")
    elif source.collected_at:
        parts.append(f"{source.collected_at} 확인")
    else:
        parts.append("기준일 미확인")
    for value in (source.domain, source.source_type, source.fact_status):
        cleaned = value.strip()
        if cleaned and cleaned not in parts:
            parts.append(cleaned)
    return " · ".join(parts)


def _source_used_sections(source: Source) -> str:
    labels: list[str] = []
    for section_id in source.used_in:
        spec = SECTION_BY_ID.get(section_id)
        label = f"{spec.display_number}장" if spec is not None else section_id.strip()
        if label and label not in labels:
            labels.append(label)
    return " · ".join(labels) or "—"


def _add_source_link(paragraph: Paragraph, source: Source) -> None:
    """자료명에 검증 가능한 원문 링크를 붙이고, 링크도 무채색으로 유지한다."""

    label = _source_label(source)
    url = source.url.strip()
    if not url.startswith(("https://", "http://")):
        run = paragraph.add_run(label)
        _set_font(run, size_pt=FONT_SIZE_TABLE_PT)
        return

    relation_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = paragraph.add_run(label)
    _set_font(
        run,
        size_pt=FONT_SIZE_TABLE_PT,
        color_rgb=COLOR_INK_RGB,
    )
    run.font.underline = True
    hyperlink.append(run._r)  # noqa: SLF001 — python-docx에 외부 링크 공개 API가 없다
    paragraph._p.append(hyperlink)  # noqa: SLF001


def _citations(report: Report) -> list[Source]:
    citations: list[Source] = []
    seen_numbers: set[int] = set()
    for item in report.citations:
        if not isinstance(item, Source) or item.number in seen_numbers:
            continue
        seen_numbers.add(item.number)
        citations.append(item)
    return citations


def _add_citations_table(doc: DocxDocument, citations: list[Source]) -> None:
    """원문을 직접 열고 위치·사용 장까지 대조할 수 있는 검증표를 낸다."""

    table = doc.add_table(rows=1, cols=5)
    table.style = TABLE_STYLE

    for idx, head in enumerate(("#", "자료", "기준일·상태", "원문 위치", "본문 사용 장")):
        cell = table.rows[0].cells[idx]
        run = cell.paragraphs[0].add_run(head)
        _set_font(run, bold=True, size_pt=FONT_SIZE_TABLE_PT)
        _shade_cell(cell, COLOR_TABLE_HEADER_HEX)

    for source in citations:
        cells = table.add_row().cells
        number_run = cells[0].paragraphs[0].add_run(str(source.number))
        _set_font(number_run, size_pt=FONT_SIZE_TABLE_PT)
        _add_source_link(cells[1].paragraphs[0], source)
        status_run = cells[2].paragraphs[0].add_run(_source_status(source))
        _set_font(status_run, size_pt=FONT_SIZE_TABLE_PT)
        location_run = cells[3].paragraphs[0].add_run(source.location.strip() or "—")
        _set_font(location_run, size_pt=FONT_SIZE_TABLE_PT)
        usage_run = cells[4].paragraphs[0].add_run(_source_used_sections(source))
        _set_font(usage_run, size_pt=FONT_SIZE_TABLE_PT)
    _set_table_geometry(table, SOURCES_TABLE_WIDTHS_DXA)


# ══════════════════════════════════════════════════════════
# 상단 — 회사명·부제·등급 라벨
# ══════════════════════════════════════════════════════════


def _cover_metadata(report: Report) -> str:
    values = [
        report.corp_type.strip(),
        f"{report.as_of_date.strip()} 기준" if report.as_of_date.strip() else "",
        report.analysis_period.strip(),
        (
            f"최신 실적 {report.latest_performance_period.strip()}"
            if report.latest_performance_period.strip()
            else ""
        ),
    ]
    return " · ".join(value for value in values if value)


def _add_subtitle(doc: DocxDocument, report: Report) -> None:
    """생성 과정 대신 회사 유형과 분석 기준만 표지에 둔다."""
    paragraph = doc.add_paragraph(style="Subtitle")
    tail_run = paragraph.add_run(_cover_metadata(report))
    _set_font(
        tail_run,
        size_pt=FONT_SIZE_SUBTITLE_PT,
        color_rgb=COLOR_MUTED_RGB,
    )


def _add_summary(doc: DocxDocument, report: Report) -> None:
    if not report.summary_items:
        return
    _add_heading(doc, "핵심 요약", level=1)
    table = doc.add_table(rows=0, cols=3)
    table.style = TABLE_STYLE
    for index, item in enumerate(report.summary_items, start=1):
        cells = table.add_row().cells
        number_run = cells[0].paragraphs[0].add_run(f"{index:02d}")
        _set_font(number_run, size_pt=FONT_SIZE_TABLE_PT, bold=True, color_rgb=(255, 255, 255))
        _shade_cell(cells[0], "171717")
        text_run = cells[1].paragraphs[0].add_run(item.text.strip())
        _set_font(text_run, size_pt=FONT_SIZE_TABLE_PT)
        spec = SECTION_BY_ID.get(item.section_id)
        related = f"{spec.display_number}장" if spec is not None else ""
        related_run = cells[2].paragraphs[0].add_run(related)
        _set_font(related_run, size_pt=FONT_SIZE_MUTED_PT, color_rgb=COLOR_MUTED_RGB)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_table_geometry(table, SUMMARY_TABLE_WIDTHS_DXA, repeat_header=False)


def _add_cover(doc: DocxDocument, report: Report) -> None:
    """회사명/보고서명과 핵심 요약을 첫 페이지의 두 명확한 영역에 배치한다."""

    top_space = doc.add_paragraph()
    top_space.paragraph_format.space_after = Pt(68)
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    company_run = title.add_run(report.company)
    _set_font(company_run, size_pt=FONT_SIZE_TITLE_PT, bold=True)
    company_run.add_break()
    report_run = title.add_run("분석 보고서")
    _set_font(report_run, size_pt=FONT_SIZE_TITLE_PT, bold=True)
    _add_subtitle(doc, report)
    summary_space = doc.add_paragraph()
    summary_space.paragraph_format.space_after = Pt(72)
    _add_summary(doc, report)
    doc.add_page_break()


# ══════════════════════════════════════════════════════════
# 본문 — 항목(칸)
# ══════════════════════════════════════════════════════════


def _section_heading(section: ReportSection) -> str:
    if section.display_number:
        tag = f"  {section.tag}" if section.tag else ""
        return f"{section.display_number}. {section.title}{tag}"
    return section_display_heading(section.cell, section.title)


def _output_sections(report: Report) -> list[ReportSection]:
    return list(report.sections)


def _add_section(doc: DocxDocument, section: ReportSection) -> None:
    _add_heading(doc, _section_heading(section), level=2)
    if not section.is_filled:
        return
    if section.prose_lines:
        _add_cited_prose(doc, section.prose_lines)
    for table in section.tables:
        _add_report_table(doc, table)


def _add_citations(doc: DocxDocument, report: Report) -> None:
    citations = _citations(report)
    if not citations:
        return

    _add_heading(doc, HEADING_SOURCES, level=2)
    _add_muted_paragraph(doc, CITATIONS_NOTE)
    _add_citations_table(doc, citations)


# ══════════════════════════════════════════════════════════
# 진입점
# ══════════════════════════════════════════════════════════


def build_docx(report: Report) -> bytes:
    """보고서 하나를 워드(.docx) 바이트로 만든다.

    canonical 공개본은 표지 → 핵심 요약 → 본문 1~9장 → 출처 순서다.
    구형 저장본은 공개 내보내지 않는다.

    Args:
        report: 06 검증을 통과한(또는 부분 통과한) 보고서 원본.

    Returns:
        `.docx` 파일 바이트. 디스크에 쓰지 않는다 — 부르는 쪽(웹)이 그대로
        내려보낸다.
    """
    report = build_published_report(report)

    doc = Document()
    _set_page_geometry(doc)
    _configure_styles(doc)
    _configure_bullet_numbering(doc)
    _clear_personal_metadata(doc, title=f"{report.company} 분석 보고서")
    _add_page_furniture(doc, report_title=f"{report.company} 분석 보고서")
    _add_cover(doc, report)

    for section in _output_sections(report):
        _add_section(doc, section)

    _add_citations(doc, report)

    _clear_revision_ids(doc)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ══════════════════════════════════════════════════════════
# 연결 지점 — 웹이 다운로드 응답을 만들 때 쓸 값
# ══════════════════════════════════════════════════════════


def build_download_filename(report: Report) -> str:
    """다운로드 파일명 — `회사명_분석_보고서_YYYY-MM-DD.docx`.

    정본: 확정/07_출력/1_흐름/01_세형태.md §워드로 받기

    ★ 생성일을 쓴다(다운로드한 «오늘»이 아니다) — 「다시 내보내기」해도 같은
      문서가 나와야 하므로(정본 §다시 내보내기) 파일명도 같아야 한다.
      기록에 생성일이 없으면(예외적) 오늘 날짜로 대신한다.

    Args:
        report: 파일명을 만들 보고서.

    Returns:
        금지 문자(`\\ / : * ? " < > |`)를 지운 파일명.
    """
    date = report.generated_at.strip() or dt.date.today().isoformat()
    company = FILENAME_FORBIDDEN_CHARS.sub("", report.company).strip()
    return f"{company or FILENAME_FALLBACK}_분석_보고서_{date}{DOCX_SUFFIX}"


def build_ascii_filename(filename: str) -> str:
    """한글 파일명에서 «안전한» ASCII 대체 이름을 만든다 (문제로그 P-77).

    ★ 왜 그냥 한글만 지우면 안 되나 — 지우고 남은 찌꺼기가 그대로 헤더에 들어간다.
      실제 사고: `루트로닉_고전압 파워 R&D 연구원_2026-08-15.docx` →
      `_  R&D _2026-08-15.docx` (밑줄로 시작·겹공백·`&`).
      이런 값은 브라우저·보안 프로그램에 따라 **헤더 전체를 무시**하게 만들고,
      그러면 파일명이 통째로 사라져 임의 이름(GUID)으로 저장된다.

    ★ 그래서 «남기는» 규칙으로 뒤집는다 — 영문·숫자·점·밑줄·붙임표만 남긴다.
      쓸 만한 글자가 없으면 순수 ASCII 자리표시자를 쓴다.
      ⚠️ 자리표시자에 한글을 쓰면 안 된다 — HTTP 헤더에 못 들어간다.

    Args:
        filename: `build_download_filename()`이 만든 한글 파일명.

    Returns:
        `.docx`로 끝나는, 순수 ASCII 파일명.
    """
    stem = filename[: -len(DOCX_SUFFIX)] if filename.endswith(DOCX_SUFFIX) else filename
    cleaned = FILENAME_ASCII_ALLOWED.sub("_", stem).strip("_.-")
    if len(cleaned) < FILENAME_ASCII_MIN_STEM:
        cleaned = FILENAME_ASCII_FALLBACK
    elif not any(char.isalpha() for char in cleaned):
        # 회사명·직무가 전부 한글이면 날짜만 남는다 — 「2026-08-15.docx」는
        # 이름 구실을 못 하므로 앞에 자리표시자를 붙인다.
        cleaned = f"{FILENAME_ASCII_FALLBACK}_{cleaned}"
    return f"{cleaned}{DOCX_SUFFIX}"


def build_content_disposition(filename: str) -> str:
    """다운로드 파일명을 `Content-Disposition` 헤더 값으로 만든다.

    ★ 왜 필요한가 — HTTP 헤더 값은 라틴-1(ASCII 계열)만 허용해 한글 파일명을
      그대로 넣을 수 없다. RFC 6266/5987의 `filename*=UTF-8''...` 확장을
      같이 넣으면, 최신 브라우저는 원래 한글 이름으로 받고 그 확장을 모르는
      옛 브라우저는 `filename=`의 ASCII 대체 이름으로 받는다(깨지지 않고
      다운로드 자체는 된다).

    Args:
        filename: `build_download_filename()`이 만든 한글 파일명.

    Returns:
        `attachment; filename="..."; filename*=UTF-8''...` 형태의 헤더 값.
    """
    encoded = urllib.parse.quote(filename)
    return (
        f'attachment; filename="{build_ascii_filename(filename)}"; '
        f"filename*=UTF-8''{encoded}"
    )


__all__ = [
    "CONTENT_TYPE_DOCX",
    "build_ascii_filename",
    "build_content_disposition",
    "build_docx",
    "build_download_filename",
]
