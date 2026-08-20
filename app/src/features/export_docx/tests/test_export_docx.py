"""워드 내보내기 시험.

★ 확인해야 하는 것 (팀장 지시) — 만든 바이트를 python-docx로 다시 열어서:
  ① 항목 개수  ② 빈칸 사유 존재  ③ 표 존재  ④ 출처 목록 존재
  ⑤ 요구역량 존재  ⑥ 파일이 안 깨지는지

  + 화면 ↔ 워드 내용 일치 관련 — 같은 제목·회사 사실·표·최종 출처를 쓰고
    근거 원문과 제작 메타문구가 반복되지 않는지 값 단위로 검증한다.
"""

from __future__ import annotations

import io
import zipfile
from dataclasses import replace

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from src.core.constants import CELL_LABELS
from src.features.export_docx.constants import (
    BODY_LINE_SPACING,
    BODY_SPACE_AFTER_PT,
    BULLET_HANGING_DXA,
    BULLET_MARKER_DXA,
    BULLET_TEXT_DXA,
    CITATIONS_NOTE,
    EMPTY_DEFAULT_REASON,
    EMPTY_PREFIX,
    FILENAME_FALLBACK,
    FONT_NAME,
    FONT_SIZE_BODY_PT,
    FONT_SIZE_HEADING_PT,
    HEADER_FOOTER_DISTANCE_IN,
    HEADING_SOURCES,
    PAGE_HEIGHT_IN,
    PAGE_MARGIN_IN,
    PAGE_WIDTH_IN,
    SOURCES_TABLE_WIDTHS_DXA,
    TABLE_CELL_BOTTOM_DXA,
    TABLE_CELL_END_DXA,
    TABLE_CELL_START_DXA,
    TABLE_CELL_TOP_DXA,
    TABLE_INDENT_DXA,
    TABLE_WIDTH_DXA,
)
from src.features.export_docx.logic import (
    _add_cited_prose,
    _add_report_table,
    build_content_disposition,
    build_docx,
    build_download_filename,
)
from src.features.pipeline.canonical_demo import build_demo_report
from src.features.pipeline.port import (
    Grade,
    Report,
    ReportSection,
    ReportTable,
    SourceStatus,
)
from src.features.report_standard import PublishBlockedError, SECTION_BY_ID

_LEGACY_SECRET = "LEGACY-JOB-POSTING-SECRET"

# ══════════════════════════════════════════════════════════
# 시험용 원본 — 실제 출고 게이트를 통과한 canonical 1~9장 보고서
# ══════════════════════════════════════════════════════════


def _make_report(**overrides) -> Report:
    report = replace(
        build_demo_report(),
        job=_LEGACY_SECRET,
        requirements=[_LEGACY_SECRET],
        sources=[SourceStatus(_LEGACY_SECRET, "failed", _LEGACY_SECRET)],
        shortfall_reasons=[_LEGACY_SECRET],
    )
    return replace(report, **overrides)


def _make_legacy_partial_report() -> Report:
    """옛 부분 보고서는 DOCX 생성 단계에서도 반드시 fail-closed여야 한다."""

    return Report(
        company="에스엠",
        job=_LEGACY_SECRET,
        corp_type="상장사",
        grade=Grade.PARTIAL,
        sections=[
            ReportSection(
                cell="1",
                title=CELL_LABELS["1"],
                lines=[("근거 계약이 없는 옛 한 장 보고서", "2")],
            )
        ],
        citations=[],
        cells={"1": True},
        shortfall_reasons=["필수 장과 canonical 근거가 없습니다"],
        generated_at="2026-08-13",
    )


def _open(data: bytes) -> Document:
    return Document(io.BytesIO(data))


def _all_text(doc: Document) -> str:
    """문단·표 안의 글자를 전부 이어붙인다 — 존재 여부를 문자열 검색으로 확인한다."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ══════════════════════════════════════════════════════════
# ① 파일이 안 깨지는가 — 바이트를 다시 열 수 있어야 한다
# ══════════════════════════════════════════════════════════


def test_바이트를_다시_열_수_있다():
    data = build_docx(_make_report())
    assert isinstance(data, bytes)
    assert len(data) > 0
    doc = _open(data)  # 예외가 나면 시험이 실패한다
    assert len(doc.paragraphs) > 0


# ══════════════════════════════════════════════════════════
# ② 항목 개수 — report.sections 하나마다 제목이 하나씩 나온다
# ══════════════════════════════════════════════════════════


def test_항목_개수가_섹션_개수와_같다():
    report = _make_report()
    doc = _open(build_docx(report))
    headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
    expected_sections = list(report.sections)
    section_headings = [
        heading
        for heading in headings
        if heading != HEADING_SOURCES
    ]
    assert len(section_headings) == len(expected_sections)
    for section in expected_sections:
        tag = f"  {section.tag}" if section.tag else ""
        assert f"{section.display_number}. {section.title}{tag}" in section_headings


def test_회사명이_문서_제목으로_들어간다():
    report = _make_report()
    doc = _open(build_docx(report))
    title = next(paragraph for paragraph in doc.paragraphs if paragraph.style.name == "Title")
    assert title.text == f"{report.company}\n분석 보고서"
    title_index = next(
        index for index, paragraph in enumerate(doc.paragraphs) if paragraph.style.name == "Title"
    )
    assert title_index < next(
        index
        for index, paragraph in enumerate(doc.paragraphs)
        if paragraph.text.startswith("1.")
    )


def test_검증된_본문만_한번_내고_근거원문을_반복하지_않는다():
    doc = Document()
    _add_cited_prose(doc, [("검증된 표시용 사업 문장입니다.", "조각 2·사업보고서")])
    text = _all_text(doc)

    assert "검증된 표시용 사업 문장입니다. 〔2〕" in text
    assert "원문 사업 문장입니다. 〔2〕" not in text
    assert "조각 2·사업보고서" not in text


# ══════════════════════════════════════════════════════════
# ③ 빈칸 사유 존재 — S6, 지우지 않는다
# ══════════════════════════════════════════════════════════


def test_legacy_부분_보고서는_DOCX를_출력하지_않는다():
    with pytest.raises(PublishBlockedError, match="canonical"):
        build_docx(_make_legacy_partial_report())


def test_필수_1에서_9장은_모두_채워진_경우에만_DOCX가_생성된다():
    report = _make_report()
    assert all(section.is_filled for section in report.sections)
    text = _all_text(_open(build_docx(report)))
    assert EMPTY_PREFIX not in text
    assert EMPTY_DEFAULT_REASON not in text


# ══════════════════════════════════════════════════════════
# ④ 표 존재 — 숫자는 문장이 아니라 워드 표로 낸다 (D13)
# ══════════════════════════════════════════════════════════


def test_숫자_표가_워드_표로_들어간다():
    report = _make_report()
    doc = _open(build_docx(report))
    expected_tables = [table for section in report.sections for table in section.tables]
    # 핵심 요약 표 + 본문 표들 + 최종 출처 표만 남는다.
    assert len(doc.tables) == 1 + len(expected_tables) + 1
    for expected_table in expected_tables:
        found = next(
            table
            for table in doc.tables
            if [cell.text for cell in table.rows[0].cells] == expected_table.headers
        )
        assert found.cell(1, 0).text == expected_table.rows[0][0]
        assert found.cell(1, 1).text == expected_table.rows[0][1]


def test_표_열_개수가_안_맞아도_안_깨진다():
    """렌더러의 방어 동작은 출고 게이트와 분리한 저수준 단위 시험으로 확인한다."""
    bad_table = ReportTable(
        caption="테스트 표",
        headers=["a", "b", "c"],
        rows=[["1", "2"]],  # 열 하나 모자람
        numeric=False,
    )
    # canonical 문서로 스타일을 먼저 구성하고, malformed 표 자체만 저수준으로 붙인다.
    doc = _open(build_docx(_make_report()))
    _add_report_table(doc, bad_table)
    assert doc.tables[-1].cell(1, 2).text == ""


# ══════════════════════════════════════════════════════════
# ⑤ 출처 목록 존재 — 자료와 기준일·상태만 간결한 표로 옮긴다
# ══════════════════════════════════════════════════════════


def test_출처_목록이_들어간다():
    report = _make_report()
    text = _all_text(_open(build_docx(report)))
    assert HEADING_SOURCES in text
    assert CITATIONS_NOTE in text
    assert "날짜까지 적었습니다" not in text
    for source in report.citations:
        assert (source.title or source.label) in text


def test_출처는_자료와_기준일_상태_표로_간결하게_낸다():
    report = _make_report()
    doc = _open(build_docx(report))
    source_table = doc.tables[-1]
    assert [source_table.cell(0, index).text for index in range(5)] == [
        "#",
        "자료",
        "기준일·상태",
        "원문 위치",
        "본문 사용 장",
    ]
    source = report.citations[0]
    row = [source_table.cell(1, index).text for index in range(5)]
    assert row[0] == str(source.number)
    assert (source.title or source.label) in row[1]
    assert source.publisher in row[1]
    assert source.disclosed_at in row[2]
    assert source.source_type in row[2]
    assert source.fact_status in row[2]
    assert row[3] == source.location
    for section_id in source.used_in:
        assert f"{SECTION_BY_ID[section_id].display_number}장" in row[4]
    assert "수집" not in _all_text(doc)


def test_출처가_없으면_빈_DOCX가_아니라_출고를_차단한다():
    report = _make_report(citations=[])
    with pytest.raises(PublishBlockedError):
        build_docx(report)


# ══════════════════════════════════════════════════════════
# 수집 현황 — P2(출력 구성 누락 0건)와 직결
# ══════════════════════════════════════════════════════════


def test_수집_과정_표는_최종_보고서에_들어가지_않는다():
    report = _make_report()
    doc = _open(build_docx(report))
    text = _all_text(doc)
    assert "어디서 가져왔나" not in text
    assert not any(table.cell(0, 0).text == "소스" for table in doc.tables)


# ══════════════════════════════════════════════════════════
# 회사분석 전용 경계 — 옛 직무/공고는 숨기고 준비 질문은 별도로 표시
# ══════════════════════════════════════════════════════════


def test_옛_직무공고_필드는_워드에_들어가지_않는다():
    report = _make_report()
    text = _all_text(_open(build_docx(report)))
    assert _LEGACY_SECRET not in text
    assert f"5. {CELL_LABELS['5']}" not in text
    assert f"6. {CELL_LABELS['6']}" not in text
    assert f"7. {CELL_LABELS['7']}" not in text
    assert f"8. {CELL_LABELS['8']}" not in text


def test_회사_사실_1에서_9장만_내고_프로그램의_작성_안내는_숨긴다():
    report = _make_report()
    text = _all_text(_open(build_docx(report)))
    for section in report.sections:
        assert section.title in text
    assert f"활용. {CELL_LABELS['활용']}" not in text
    assert "확인된 회사 사실" not in text
    assert "프로그램이 제안" not in text
    assert "활용 질문" not in text


def test_legacy_5번_한_장_기록은_본문을_내지_않고_차단한다():
    report = replace(
        _make_legacy_partial_report(),
        sections=[
            ReportSection(cell="5", title=CELL_LABELS["5"], lines=[(_LEGACY_SECRET, "")])
        ],
    )
    with pytest.raises(PublishBlockedError):
        build_docx(report)


# ══════════════════════════════════════════════════════════
# 등급 라벨 — 완성이면 안 띄운다 (result.html과 같은 조건)
# ══════════════════════════════════════════════════════════


def test_완성도와_내부_검증_문구는_최종_보고서에_들어가지_않는다():
    for grade in (Grade.COMPLETE, Grade.PARTIAL, Grade.INCOMPLETE):
        report = _make_report(grade=grade)
        text = _all_text(_open(build_docx(report)))
        assert "🟡" not in text
        assert "🔴" not in text
        assert report.shortfall_reasons[0] not in text


# ══════════════════════════════════════════════════════════
# 제출용 문서 프리셋 — 내용은 그대로 두고 구조·가독성만 검증
# ══════════════════════════════════════════════════════════


def test_business_brief_페이지와_본문_스타일을_명시한다():
    doc = _open(build_docx(_make_report()))
    section = doc.sections[0]
    assert section.page_width == Inches(PAGE_WIDTH_IN)
    assert section.page_height == Inches(PAGE_HEIGHT_IN)
    assert section.top_margin == Inches(PAGE_MARGIN_IN)
    assert section.right_margin == Inches(PAGE_MARGIN_IN)
    assert section.bottom_margin == Inches(PAGE_MARGIN_IN)
    assert section.left_margin == Inches(PAGE_MARGIN_IN)
    expected_distance_twips = round(HEADER_FOOTER_DISTANCE_IN * 1440)
    assert section.header_distance.twips == expected_distance_twips
    assert section.footer_distance.twips == expected_distance_twips

    normal = doc.styles["Normal"]
    assert normal.font.name == FONT_NAME
    assert normal.font.size == Pt(FONT_SIZE_BODY_PT)
    assert normal.paragraph_format.space_after == Pt(BODY_SPACE_AFTER_PT)
    assert normal.paragraph_format.line_spacing == BODY_LINE_SPACING

    heading = doc.styles["Heading 2"]
    assert heading.font.name == FONT_NAME
    assert heading.font.size == Pt(FONT_SIZE_HEADING_PT)
    assert heading.paragraph_format.keep_with_next is True


def test_글머리는_실제_numbering과_정확한_들여쓰기를_쓴다():
    doc = _open(build_docx(_make_report()))
    style = doc.styles["List Bullet"]
    assert style.paragraph_format.left_indent == Inches(0.5)
    assert style.paragraph_format.first_line_indent == Inches(-0.25)

    num_id = str(style.element.pPr.numPr.numId.val)
    numbering = doc.part.numbering_part.element
    num = numbering.xpath(f"./w:num[@w:numId='{num_id}']")[0]
    abstract_id = str(num.abstractNumId.val)
    level = numbering.xpath(
        f"./w:abstractNum[@w:abstractNumId='{abstract_id}']/w:lvl[@w:ilvl='0']"
    )[0]
    indent = level.find("w:pPr/w:ind", level.nsmap)
    tab = level.find("w:pPr/w:tabs/w:tab", level.nsmap)
    assert indent.get(qn("w:left")) == str(BULLET_TEXT_DXA)
    assert indent.get(qn("w:hanging")) == str(BULLET_HANGING_DXA)
    assert tab.get(qn("w:pos")) == str(BULLET_MARKER_DXA)


def test_모든_표가_고정_DXA폭과_셀여백을_쓴다():
    doc = _open(build_docx(_make_report()))
    for index, table in enumerate(doc.tables):
        table_properties = table._tbl.tblPr
        assert table_properties.find(qn("w:tblW")).get(qn("w:w")) == str(
            TABLE_WIDTH_DXA
        )
        assert table_properties.find(qn("w:tblInd")).get(qn("w:w")) == str(
            TABLE_INDENT_DXA
        )
        assert table.autofit is False
        assert table_properties.find(qn("w:tblLayout")).get(qn("w:type")) == "fixed"

        grid_widths = tuple(
            int(column.get(qn("w:w")))
            for column in table._tbl.tblGrid.findall(qn("w:gridCol"))
        )
        assert len(grid_widths) == len(table.columns)
        assert sum(grid_widths) == TABLE_WIDTH_DXA
        for row in table.rows:
            for cell, width in zip(row.cells, grid_widths, strict=True):
                assert cell._tc.tcPr.tcW.get(qn("w:w")) == str(width)
        assert not table._tbl.xpath(".//w:trHeight"), "고정 행 높이는 글자를 자를 수 있다"
        # 핵심 요약은 각 행이 독립 항목이라 머리글이 없다. 본문·출처 표만 반복한다.
        if index > 0:
            assert table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None

        margins = table_properties.find(qn("w:tblCellMar"))
        expected = {
            "top": TABLE_CELL_TOP_DXA,
            "start": TABLE_CELL_START_DXA,
            "bottom": TABLE_CELL_BOTTOM_DXA,
            "end": TABLE_CELL_END_DXA,
        }
        for side, width in expected.items():
            assert margins.find(qn(f"w:{side}")).get(qn("w:w")) == str(width)

    source_grid = tuple(
        int(column.get(qn("w:w")))
        for column in doc.tables[-1]._tbl.tblGrid.findall(qn("w:gridCol"))
    )
    assert source_grid == SOURCES_TABLE_WIDTHS_DXA


def test_공개_DOCX에는_빈칸_callout이_존재하지_않는다():
    report = _make_report()
    doc = _open(build_docx(report))
    assert all(section.is_filled for section in report.sections)
    assert not any(EMPTY_PREFIX in paragraph.text for paragraph in doc.paragraphs)
    assert not any(EMPTY_DEFAULT_REASON in paragraph.text for paragraph in doc.paragraphs)


def test_표_속성이_OOXML_스키마_순서를_지킨다():
    doc = _open(build_docx(_make_report()))
    tblpr_order = (
        "tblStyle",
        "tblW",
        "jc",
        "tblCellSpacing",
        "tblInd",
        "tblBorders",
        "shd",
        "tblLayout",
        "tblCellMar",
        "tblLook",
    )
    for table in doc.tables:
        names = [child.tag.rsplit("}", 1)[-1] for child in table._tbl.tblPr]
        present = [name for name in names if name in tblpr_order]
        assert present == [name for name in tblpr_order if name in present]
        required = ("tblW", "tblInd", "tblBorders", "tblLayout", "tblCellMar")
        assert all(name in names for name in required)


def test_출처는_본문_목록이_아닌_고정폭_검증표로_낸다():
    doc = _open(build_docx(_make_report()))
    source_table = doc.tables[-1]
    grid = tuple(
        int(column.get(qn("w:w")))
        for column in source_table._tbl.tblGrid.findall(qn("w:gridCol"))
    )

    assert grid == SOURCES_TABLE_WIDTHS_DXA
    assert source_table.autofit is False
    assert not any(paragraph.style.name == "Report Source" for paragraph in doc.paragraphs)


def test_머리말_페이지번호와_개인_메타데이터_없음을_확인한다():
    report = _make_report()
    data = build_docx(report)
    doc = _open(data)
    section = doc.sections[0]
    assert section.different_first_page_header_footer is True
    assert section.first_page_header.paragraphs[0].text == ""
    assert section.first_page_footer.paragraphs[0].text == ""
    assert section.header.paragraphs[0].text == f"{report.company} 분석 보고서"
    assert "PAGE" in section.footer._element.xml
    assert section.footer.paragraphs[0].text.startswith("페이지 ")
    assert doc.core_properties.author in (None, "")
    assert doc.core_properties.last_modified_by in (None, "")
    assert doc.core_properties.comments in (None, "")
    assert doc.core_properties.title == f"{report.company} 분석 보고서"

    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        assert "docProps/custom.xml" not in archive.namelist()
        core_xml = archive.read("docProps/core.xml").decode("utf-8")
        assert f"{report.company} 분석 보고서" in core_xml
        story_parts = [
            name
            for name in archive.namelist()
            if name == "word/document.xml"
            or name.startswith("word/header")
            or name.startswith("word/footer")
        ]
        assert all(b"w:rsid" not in archive.read(name) for name in story_parts)


# ══════════════════════════════════════════════════════════
# 파일명 규칙 — 확정/07_출력/1_흐름/01_세형태.md §워드로 받기
# ══════════════════════════════════════════════════════════


def test_파일명_형식():
    report = _make_report(company="에스엠", job="마케팅", generated_at="2026-08-15")
    assert build_download_filename(report) == "에스엠_분석_보고서_2026-08-15.docx"


def test_파일명_금지문자를_지운다():
    report = _make_report(company='나*쁜/회:사?"<>|', job="영업", generated_at="2026-08-15")
    filename = build_download_filename(report)
    assert filename == "나쁜회사_분석_보고서_2026-08-15.docx"
    for ch in '\\/:*?"<>|':
        assert ch not in filename


def test_생성일이_없으면_오늘_날짜로_대신한다():
    report = _make_report(generated_at="")
    filename = build_download_filename(report)
    assert filename.endswith(".docx")
    # YYYY-MM-DD 10글자 날짜가 붙어 있어야 한다 (오늘 날짜는 시험 시각마다 달라 형식만 본다).
    date_part = filename[:-5].rsplit("_", 1)[-1]
    assert len(date_part) == 10 and date_part.count("-") == 2


def test_회사명이_전부_금지문자면_자리표시자를_쓴다():
    report = _make_report(company="///", job="영업", generated_at="2026-08-15")
    filename = build_download_filename(report)
    assert filename == f"{FILENAME_FALLBACK}_분석_보고서_2026-08-15.docx"


# ══════════════════════════════════════════════════════════
# 다운로드 헤더 — 한글 파일명이 깨지지 않는 법 (RFC 5987)
# ══════════════════════════════════════════════════════════


def test_content_disposition_한글_파일명():
    header = build_content_disposition("에스엠_마케팅_2026-08-15.docx")
    assert header.startswith("attachment; filename=")
    assert "filename*=UTF-8''" in header
    # ASCII 대체 이름에는 한글이 없어야 한다(오래된 클라이언트가 이 부분만 읽는다).
    ascii_part = header.split('filename="')[1].split('"')[0]
    assert all(ord(c) < 128 for c in ascii_part)
    # UTF-8 부분을 다시 풀면 원래 파일명이 나와야 한다.
    import urllib.parse

    encoded = header.split("UTF-8''")[1]
    assert urllib.parse.unquote(encoded) == "에스엠_마케팅_2026-08-15.docx"


def test_content_disposition_전부_한글이면_대체_이름을_쓴다():
    header = build_content_disposition("에스엠_마케팅.docx")
    ascii_part = header.split('filename="')[1].split('"')[0]
    assert ascii_part  # 비어 있지 않다 — 빈 filename은 일부 클라이언트에서 깨진다
